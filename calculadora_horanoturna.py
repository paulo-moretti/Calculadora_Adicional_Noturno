import tkinter as tk
from tkinter import filedialog, messagebox
from datetime import datetime, timedelta
from docx import Document

# Definição do período noturno
HORARIO_NOTURNO_INICIO = 22 * 60  # 22:00 em minutos (1320)
HORARIO_NOTURNO_FIM = 5 * 60       # 05:00 em minutos (300)

# Lista para armazenar os dados antes de gerar o documento
dados = []

def converter_horas(tempo):
    try:
        h, m = map(int, tempo.split(":"))
        return h * 60 + m
    except ValueError:
        return None

def calcular_horas(entrada, intervalo, retorno, saida):
    if not entrada or not intervalo or not retorno or not saida:
        return None  # Ignorar dias sem horários preenchidos

    entrada_min = converter_horas(entrada)
    intervalo_min = converter_horas(intervalo)
    retorno_min = converter_horas(retorno)
    saida_min = converter_horas(saida)

    if None in [entrada_min, intervalo_min, retorno_min, saida_min]:
        return None

    if intervalo_min < entrada_min:
        intervalo_min += 1440
    if retorno_min < intervalo_min:
        retorno_min += 1440
    if saida_min < retorno_min:
        saida_min += 1440

    total_trabalhado = (intervalo_min - entrada_min) + (saida_min - retorno_min)

    def calcular_horas_noturnas(inicio, fim):
        minutos_noturnos = 0
        if fim < inicio:
            fim += 1440
        for minuto in range(inicio, fim):
            minuto_do_dia = minuto % 1440
            if HORARIO_NOTURNO_INICIO <= minuto_do_dia or minuto_do_dia < HORARIO_NOTURNO_FIM:
                minutos_noturnos += 1
        return minutos_noturnos

    horas_noturnas_reais = (calcular_horas_noturnas(entrada_min, intervalo_min) +
                             calcular_horas_noturnas(retorno_min, saida_min))
    horas_noturnas_computadas = (horas_noturnas_reais * 60) / 52.5

    def formatar_horas(minutos):
        horas = max(0, int(minutos // 60))
        minutos_restantes = max(0, int(minutos % 60))
        return f"{horas:02d}:{minutos_restantes:02d}"

    return {
        "Horas Normais": formatar_horas(total_trabalhado),
        "Horas Noturnas Computadas": formatar_horas(horas_noturnas_computadas)
    }

def processar_dados():
    global dados
    dados = []
    text_resultado.delete("1.0", tk.END)
    for i in range(31):
        data = entry_datas[i].get()
        entrada = entry_entradas[i].get()
        intervalo_inicio = entry_intervalos_ini[i].get()
        intervalo_fim = entry_intervalos_fim[i].get()
        saida = entry_saidas[i].get()

        resultado = calcular_horas(entrada, intervalo_inicio, intervalo_fim, saida)

        if resultado:
            dados.append([data, entrada, intervalo_inicio, intervalo_fim, saida,
                          resultado["Horas Normais"], resultado["Horas Noturnas Computadas"]])
        else:
            dados.append([data, "***", "***", "***", "***", "***", "***"])

    for linha in dados:
        text_resultado.insert(tk.END,
            f"Data: {linha[0]} | Horas Normais: {linha[5]} | Horas Noturnas Computadas: {linha[6]}\n")

    btn_download.config(state=tk.NORMAL)

def gerar_documento():
    doc = Document()
    doc.add_heading('Relatório de Cálculo de Horas Trabalhadas', level=1)
    tabela = doc.add_table(rows=1, cols=7)
    tabela.style = 'Table Grid'
    cabecalhos = ["Data", "Entrada", "Intervalo Início", "Intervalo Fim", "Saída", "Horas Normais", "Horas Noturnas Computadas"]
    hdr_cells = tabela.rows[0].cells
    for i, texto in enumerate(cabecalhos):
        hdr_cells[i].text = texto

    # Adiciona as linhas com os dados diários
    for linha in dados:
        row_cells = tabela.add_row().cells
        for i, valor in enumerate(linha):
            row_cells[i].text = str(valor)
    
    # Função auxiliar para converter "HH:MM" para minutos
    def time_to_minutes(t):
        try:
            h, m = map(int, t.split(":"))
            return h * 60 + m
        except Exception:
            return 0

    # Soma os totais apenas para linhas com dados válidos (evita "***")
    total_normais = 0
    total_noturnas = 0
    for linha in dados:
        if linha[5] != "***":
            total_normais += time_to_minutes(linha[5])
        if linha[6] != "***":
            total_noturnas += time_to_minutes(linha[6])
    
    # Função auxiliar para converter minutos para "HH:MM"
    def minutes_to_time(m):
        h = m // 60
        m = m % 60
        return f"{h:02d}:{m:02d}"
    
    total_normais_str = minutes_to_time(total_normais)
    total_noturnas_str = minutes_to_time(total_noturnas)

    # Adiciona a linha de totais
    totals_row = tabela.add_row().cells
    totals_row[0].text = "Totais"  # Você pode mesclar células se desejar um efeito diferente
    totals_row[1].text = ""
    totals_row[2].text = ""
    totals_row[3].text = ""
    totals_row[4].text = ""
    totals_row[5].text = total_normais_str
    totals_row[6].text = total_noturnas_str

    # Aplica a cor de fundo amarela na linha de totais
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    shading_elm = parse_xml(r'<w:shd {} w:fill="FFFF00"/>'.format(nsdecls('w')))
    for cell in totals_row:
        cell._tc.get_or_add_tcPr().append(shading_elm)

    caminho_arquivo = filedialog.asksaveasfilename(
        defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
    if caminho_arquivo:
        doc.save(caminho_arquivo)
        messagebox.showinfo("Sucesso", f"Relatório salvo com sucesso: {caminho_arquivo}")

# Configuração da janela principal
root = tk.Tk()
root.title("Calculadora de Horas Trabalhadas")
root.geometry("1200x800")
root.configure(bg="#F7F7F7")

# Cabeçalho
header = tk.Label(root, text="Calculadora de Horas Trabalhadas", bg="#F7F7F7", fg="#2C3E50")
header.pack(pady=10)

# Container para a área de entradas com scroll
container = tk.Frame(root, bg="#F7F7F7")
container.pack(fill="both", expand=True, padx=20, pady=10)

canvas = tk.Canvas(container, bg="#F7F7F7", highlightthickness=0)
canvas.pack(side="left", fill="both", expand=True)

scrollbar = tk.Scrollbar(container, orient="vertical", command=canvas.yview)
scrollbar.pack(side="right", fill="y")
canvas.configure(yscrollcommand=scrollbar.set)

entry_frame = tk.Frame(canvas, bg="#F7F7F7")
canvas.create_window((0, 0), window=entry_frame, anchor="nw")

def on_configure(event):
    canvas.configure(scrollregion=canvas.bbox("all"))
entry_frame.bind("<Configure>", on_configure)

# Cabeçalhos da tabela de entradas
colunas = ["Data", "Entrada", "Intervalo Início", "Intervalo Fim", "Saída"]
for j, texto in enumerate(colunas):
    lbl = tk.Label(entry_frame, text=texto, bg="#AED6F1", fg="#2C3E50", padx=5, pady=5, borderwidth=1, relief="raised")
    lbl.grid(row=0, column=j, sticky="nsew", padx=1, pady=1)

entry_datas = []
entry_entradas = []
entry_intervalos_ini = []
entry_intervalos_fim = []
entry_saidas = []

for i in range(31):
    ent_date = tk.Entry(entry_frame, width=12)
    ent_date.grid(row=i+1, column=0, padx=1, pady=1)
    entry_datas.append(ent_date)

    ent_entrada = tk.Entry(entry_frame, width=12)
    ent_entrada.grid(row=i+1, column=1, padx=1, pady=1)
    entry_entradas.append(ent_entrada)

    ent_intervalo_ini = tk.Entry(entry_frame, width=12)
    ent_intervalo_ini.grid(row=i+1, column=2, padx=1, pady=1)
    entry_intervalos_ini.append(ent_intervalo_ini)

    ent_intervalo_fim = tk.Entry(entry_frame, width=12)
    ent_intervalo_fim.grid(row=i+1, column=3, padx=1, pady=1)
    entry_intervalos_fim.append(ent_intervalo_fim)

    ent_saida = tk.Entry(entry_frame, width=12)
    ent_saida.grid(row=i+1, column=4, padx=1, pady=1)
    entry_saidas.append(ent_saida)

# Área inferior: botões e resultados
bottom_frame = tk.Frame(root, bg="#F7F7F7")
bottom_frame.pack(fill="x", padx=20, pady=10)

btn_calcular = tk.Button(bottom_frame, text="Calcular", command=processar_dados, bg="#27AE60", fg="white", padx=10, pady=5)
btn_calcular.pack(side="left", padx=10)

btn_download = tk.Button(bottom_frame, text="Baixar Relatório Word", command=gerar_documento, state=tk.DISABLED, bg="#1A5276", fg="white", padx=10, pady=5)
btn_download.pack(side="left", padx=10)

# Área de resultados com scroll
result_container = tk.Frame(root, bg="#F7F7F7")
result_container.pack(fill="both", expand=True, padx=20, pady=10)

text_resultado = tk.Text(result_container, height=10, bg="white", fg="#2C3E50", borderwidth=2, relief="groove")
text_resultado.pack(side="left", fill="both", expand=True)

result_scrollbar = tk.Scrollbar(result_container, command=text_resultado.yview)
result_scrollbar.pack(side="right", fill="y")
text_resultado.config(yscrollcommand=result_scrollbar.set)

root.mainloop()
